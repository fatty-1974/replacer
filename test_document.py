from docx import Document

# Создаем тестовый документ
doc = Document()

# Добавляем заголовок
doc.add_heading('Тестовый документ', 0)

# Добавляем параграфы с текстом для замены
paragraph1 = doc.add_paragraph('Это тестовый документ для проверки замены текста.')
paragraph2 = doc.add_paragraph('В этом документе мы будем заменять слово "тестовый" на другое слово.')
paragraph3 = doc.add_paragraph('Также проверим замену в таблице ниже:')

# Создаем таблицу
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'

# Заполняем таблицу
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Исходный текст'
hdr_cells[1].text = 'Новый текст'

row_cells = table.rows[1].cells
row_cells[0].text = 'тестовый'
row_cells[1].text = 'замененный'

row_cells = table.rows[2].cells
row_cells[0].text = 'Это тестовый пример'
row_cells[1].text = 'Это замененный пример'

# Сохраняем документ
doc.save('test_files/test_document.docx')
print('Тестовый документ создан в test_files/test_document.docx')