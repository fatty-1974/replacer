from docx_processor_fixed import DocxProcessor
import os

# Создаем тестовую папку если её нет
if not os.path.exists("test_files"):
    os.makedirs("test_files")

# Создаем тестовый документ если его нет
test_file_path = "test_files/test_document.docx"
if not os.path.exists(test_file_path):
    from docx import Document
    doc = Document()
    doc.add_heading('Тестовый документ', 0)
    doc.add_paragraph('Это тестовый документ для проверки замены текста.')
    doc.add_paragraph('В этом документе мы будем заменять слово "тестовый" на другое слово.')
    doc.save(test_file_path)
    print('Тестовый документ создан')

# Создаем процессор и тестируем замену текста
processor = DocxProcessor()

# Тестируем замену текста
print("Тестируем замену текста...")
files_processed = processor.process_folder("test_files", "тестовый", "ПРОВЕРОЧНЫЙ")
print(f"Обработано файлов: {files_processed}")

print("Тест завершен. Проверьте файл test_files/test_document.docx")