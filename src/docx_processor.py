import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocxProcessor:
    def process_folder(self, folder_path, search_text, replace_text):
        files_processed = 0
        for filename in os.listdir(folder_path):
            if filename.endswith('.docx') and not filename.startswith('~$'):
                file_path = os.path.join(folder_path, filename)
                try:
                    self.process_file(file_path, search_text, replace_text)
                    files_processed += 1
                except Exception as e:
                    print(f"Ошибка обработки файла {filename}: {e}")
        return files_processed

    def process_file(self, file_path, search_text, replace_text):
        doc = Document(file_path)
        self.replace_block(doc, search_text, replace_text)
        doc.save(file_path)

    def normalize(self, text):
        """Удаляет пробелы, табы, переносы и приводит к нижнему регистру"""
        return re.sub(r'[\s\r\n\t\x07]+', '', text or '').lower()

    def replace_block(self, doc, search_text, replace_text):
        """Заменяет группу абзацев, если их текст совпадает без пробелов"""
        norm_search = self.normalize(search_text)
        if not norm_search:
            return

        # Собираем все абзацы документа в одну строку для поиска
        all_text = ''.join(p.text for p in doc.paragraphs)
        norm_all = self.normalize(all_text)
        if norm_search not in norm_all:
            return

        # Находим стартовый абзац
        joined_search = re.sub(r'[\r\n]+', '', search_text).strip()
        start_idx = None
        for i, p in enumerate(doc.paragraphs):
            if self.normalize(p.text).startswith(self.normalize(joined_search[:30])):
                start_idx = i
                break
        if start_idx is None:
            return

        # Определяем конец блока
        end_idx = start_idx + 1
        found = False
        while end_idx <= len(doc.paragraphs):
            block_text = ''.join(p.text for p in doc.paragraphs[start_idx:end_idx])
            if self.normalize(block_text).find(norm_search) != -1:
                found = True
                break
            end_idx += 1
        if not found:
            return

        # Сохраняем форматирование первого абзаца
        first_para = doc.paragraphs[start_idx]
        first_run = first_para.runs[0] if first_para.runs else None
        alignment = first_para.alignment

        # Определяем место вставки
        insert_before = doc.paragraphs[end_idx] if end_idx < len(doc.paragraphs) else None

        # Удаляем старые абзацы
        for k in range(end_idx - 1, start_idx - 1, -1):
            p = doc.paragraphs[k]
            p._element.getparent().remove(p._element)

        # Вставляем новые абзацы в правильном порядке
        new_lines = [line for line in replace_text.split('\n')]
        for line in new_lines:
            text = line.strip()
            if insert_before is not None:
                new_p = insert_before.insert_paragraph_before(text)
            else:
                new_p = doc.add_paragraph(text)
            new_p.alignment = alignment  # сохраняем выравнивание
            if first_run:
                if not new_p.runs:
                    r = new_p.add_run(text)
                    self.copy_formatting(r, first_run)
                else:
                    self.copy_formatting(new_p.runs[0], first_run)

    def copy_formatting(self, target_run, source_run):
        """Копирует базовое форматирование"""
        if not source_run:
            return
        try:
            if source_run.font.name:
                target_run.font.name = source_run.font.name
        except Exception:
            pass
        try:
            if source_run.font.size:
                target_run.font.size = source_run.font.size
        except Exception:
            pass
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        try:
            if source_run.font.color and getattr(source_run.font.color, 'rgb', None):
                target_run.font.color.rgb = source_run.font.color.rgb
        except Exception:
            pass
